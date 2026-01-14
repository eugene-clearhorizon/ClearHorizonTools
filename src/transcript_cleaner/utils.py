import os
import webvtt
from docx import Document

def format_duration(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    return f"{hours:02}:{minutes:02}:{seconds:02}"

def process_vtt_to_docx(file, output_filename):
    doc = Document()
    doc.add_heading(f"Transcript for {os.path.basename(file)}", level=1)
    
    vtt = webvtt.read(file)

    interview_length_seconds = vtt[-1].end_in_seconds
    interview_length = format_duration(interview_length_seconds)
    
    speaker_names = set(n.voice if n.voice else "Unknown Speaker" for n in vtt)
    
    doc.add_paragraph(f"Interview length: {interview_length}")
    doc.add_paragraph(f"Speakers: {'; '.join(sorted(speaker_names))}")
    doc.add_paragraph("")
    
    unwanted_words = {'', 'Hmm. Mm hmm mm hmm.','Mm-hmm. Mm-hmm.','Mm-hmm. Hmm.', 'Mm hmm mm hmm.', 'Mm hmm. Mm hmm.', 'Mm hmm mm hmm. Mm hmm.', 'Uh-huh.', 'Yeah. And so.', 'OK. Yeah, yeah.', 'Umm.', 'Yeah, yeah.', 'Awesome.', 
                       'Mm-hmm.', 'Mm-hmm. Mm-hmm.',  'OK.', 'Right.', 'Right?', 'OK. Yeah.', 'So.', 'Uh.', 'Hmm.', 'Hmm yeah.', 'OK, cool.', 
                      'Yeah, cool.', 'Ohh.', 'Um', 'Um?', 'Um.', 'Umm?', 'Cool.', 'Mm-hmm.', 'Mm hmm.', 'Huh.', 
                      'Ohh uh-huh.', 'Hmm. MMM.', 'Yeah. Wow.', 'Ohh wow wow.', 'Um.', 'Uh huh.', 'MMM.', 'K.', 'Oh. OK.', 'Oh. OK. '}
    
    last_speaker = None
    paragraph = doc.add_paragraph()
    combined_text = ""

    for i, n in enumerate(vtt):
        speaker_name = n.voice if n.voice else "Unknown Speaker"
        text = n.text.replace("\n", " ").strip()
        
        is_unwanted_isolated = (
            text in unwanted_words and 
            (i == 0 or vtt[i - 1].voice != speaker_name) and 
            (i == len(vtt) - 1 or vtt[i + 1].voice != speaker_name)
        )
        
        if is_unwanted_isolated:
            continue

        if speaker_name != last_speaker:
            if combined_text:
                paragraph.add_run(combined_text + " ")
                combined_text = ""
            if last_speaker is not None:
                paragraph.add_run("\n\n")
            paragraph.add_run(f"{speaker_name}: ").bold = True
            last_speaker = speaker_name
        
        combined_text += text + " "
    
    if combined_text:
        paragraph.add_run(combined_text.strip() + " ")

    doc.save(output_filename)
